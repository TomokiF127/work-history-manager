#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import json

def read_skill_sheet(filepath):
    """スキルシートの構造を読み取って分析"""
    doc = Document(filepath)
    
    content = {
        "tables": [],
        "paragraphs": []
    }
    
    # テーブルの内容を抽出
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        content["tables"].append({
            "index": i,
            "data": table_data
        })
    
    # パラグラフの内容を抽出
    for para in doc.paragraphs:
        if para.text.strip():
            content["paragraphs"].append(para.text.strip())
    
    return content

if __name__ == "__main__":
    filepath = "/Users/tomo/EG/スキルシート/【F・T】スキルシート.docx"
    content = read_skill_sheet(filepath)
    
    print("=== スキルシートの構造 ===\n")
    
    print(f"テーブル数: {len(content['tables'])}")
    print(f"パラグラフ数: {len(content['paragraphs'])}\n")
    
    # 各テーブルの詳細を表示
    for i, table in enumerate(content['tables']):
        print(f"\n--- テーブル {i+1} ---")
        print(f"行数: {len(table['data'])}")
        print(f"列数: {len(table['data'][0]) if table['data'] else 0}")
        
        # 全ての行を表示（構造把握のため）
        for j, row in enumerate(table['data']):
            if j == 0:  # ヘッダー行
                print(f"ヘッダー: {row}")
            elif j < 3:  # 最初の数行のみ詳細表示
                print(f"行{j}: {[cell[:30] + '...' if len(cell) > 30 else cell for cell in row]}")
            elif j == len(table['data']) - 1:  # 最後の行
                print(f"最終行: {[cell[:30] + '...' if len(cell) > 30 else cell for cell in row]}")
        
        if len(table['data']) > 4:
            print(f"... (中間 {len(table['data']) - 4} 行省略)")
        print(f"テーブル {i+1} 構造: {len(table['data'])}行 x {len(table['data'][0]) if table['data'] else 0}列")
    
    # パラグラフの概要
    print("\n--- パラグラフ ---")
    for para in content['paragraphs'][:10]:
        print(f"- {para[:50]}..." if len(para) > 50 else f"- {para}")