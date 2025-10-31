#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
スキルシート出力サービス
"""

from typing import List, Dict, Any, Optional
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from sqlalchemy.orm import Session

from services.repository import Repository
from services.stats import StatsService
from models import Project, TechUsage, SelfPR


class SkillSheetExportService:
    """スキルシートエクスポートサービス"""
    
    def __init__(self, session: Session):
        self.session = session
        self.repo = Repository(session)
        self.stats_service = StatsService(session)
    
    def _set_table_borders(self, table):
        """テーブルに罫線を設定（縦罫線は破線）"""
        tbl = table._tbl
        tbl_borders = OxmlElement('w:tblBorders')
        
        # 外枠は実線、縦罫線は破線に設定
        border_settings = {
            'top': 'single',
            'left': 'single', 
            'bottom': 'single',
            'right': 'single',
            'insideH': 'single',  # 横線は実線
            'insideV': 'dashed'   # 縦線は破線
        }
        
        for border_name, border_style in border_settings.items():
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), '4')  # 0.5pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tbl_borders.append(border)
        
        tbl_pr = tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        tbl_pr.append(tbl_borders)
    
    def _set_cell_properties(self, cell, width_cm=None, vertical_align='top'):
        """セルのプロパティを設定"""
        # セル幅設定
        if width_cm:
            cell.width = Cm(width_cm)
        
        # 垂直方向の配置
        cell.vertical_alignment = getattr(WD_ALIGN_VERTICAL, vertical_align.upper())
        
        # セル内のパラグラフの書式設定
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # フォント設定
            for run in paragraph.runs:
                run.font.name = 'MS 明朝'
                run.font.size = Pt(9)
    
    def _format_header_cell(self, cell, text):
        """ヘッダーセルの書式設定"""
        cell.text = text
        # 背景色を薄いグレーに設定
        shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # フォント設定
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'MS 明朝'
                run.font.size = Pt(9)
    
    def generate_skill_sheet_data(self, name: str = "氏名") -> Dict[str, Any]:
        """スキルシートデータを生成
        
        Args:
            name: 氏名
        
        Returns:
            スキルシートデータ
        """
        data = {
            "header": {
                "title": "職務経歴書",
                "date": datetime.now().strftime("%Y年%m月現在"),
                "name": name
            },
            "projects": self._generate_projects_data(),
            "technical_skills": self._generate_technical_skills_data(),
            "qualifications": self._generate_qualification_data(),
            "other_experiences": self._generate_other_experience_data(),
            "self_pr": self._generate_self_pr_data()
        }
        
        return data
    
    def _generate_projects_data(self) -> List[Dict[str, Any]]:
        """プロジェクトデータを生成（契約会社で合算）"""
        projects = self.repo.get_all_projects()
        
        # 契約会社ごとにグループ化
        company_groups = defaultdict(list)
        for project in projects:
            company = project.contract_company or "未設定"
            company_groups[company].append(project)
        
        result = []
        
        for company, company_projects in company_groups.items():
            # 契約会社内でプロジェクトを開始日順にソート
            company_projects.sort(key=lambda p: p.project_start or "9999-99-99")
            
            # 参画期間を計算（同じ契約会社の最初と最後）
            first_start = None
            last_end = None
            
            for project in company_projects:
                if project.project_start:
                    if not first_start or project.project_start < first_start:
                        first_start = project.project_start
                
                if project.project_end:
                    if not last_end or project.project_end > last_end:
                        last_end = project.project_end
                elif not last_end:
                    last_end = None  # 現在進行中
            
            # 各プロジェクトをエクスポート
            for project in company_projects:
                project_data = {
                    "participation_period": self._format_period_with_duration(first_start, last_end),
                    "project_period": self._format_period(project.project_start, project.project_end),
                    "project_name": project.name,
                    "business_content": project.work_summary,
                    "project_detail": project.detail,
                    "environment": self._get_project_environment(project.id),
                    "role": self._get_project_roles_sorted(project.id),
                    "task": self._get_project_tasks_sorted(project.id),
                    "team_size": project.scale_text or "",
                    "contract_company": company,
                    "end_user": project.end_user or "",
                    "_sort_key": project.project_start or "1900-01-01"  # ソート用キーを追加
                }
                result.append(project_data)
        
        # プロジェクト開始日の新しい順にソート（新しい案件が最初に来る）
        result.sort(key=lambda p: p.get("_sort_key", "1900-01-01"), reverse=True)
        
        return result
    
    def _get_project_environment(self, project_id: int) -> Dict[str, List[str]]:
        """プロジェクトの開発環境を取得"""
        environment = {
            "os": [],
            "language": [],
            "db": [],
            "framework": [],
            "tool": [],
            "cloud": []
        }
        
        # 各技術カテゴリの情報を取得
        for kind in ["os", "language", "framework", "tool", "cloud", "db"]:
            tech_ids = self.repo.get_project_techs(project_id, kind)
            techs = self.repo.get_master_by_kind(kind)
            
            for tech in techs:
                if tech.id in tech_ids:
                    if kind == "framework":
                        environment["framework"].append(tech.name)
                    else:
                        environment[kind].append(tech.name)
        
        return environment
    
    def _get_project_roles_sorted(self, project_id: int) -> str:
        """プロジェクトの役割を順序で取得"""
        from models import ProjectRole, Role
        
        # プロジェクトに関連する役割を順序で取得
        roles = self.session.query(Role).join(ProjectRole).filter(
            ProjectRole.project_id == project_id
        ).order_by(Role.order_index, Role.name).all()
        
        # 単一プロジェクトの場合は従来通り
        project = self.repo.get_project_by_id(project_id)
        if project and project.role and not roles:
            return project.role.name
        
        return ", ".join([role.name for role in roles]) if roles else ""
    
    def _get_project_tasks_sorted(self, project_id: int) -> str:
        """プロジェクトの作業を順序で取得"""
        from models import ProjectTask, Task
        
        # プロジェクトに関連する作業を順序で取得
        tasks = self.session.query(Task).join(ProjectTask).filter(
            ProjectTask.project_id == project_id
        ).order_by(Task.order_index, Task.name).all()
        
        # 単一プロジェクトの場合は従来通り
        project = self.repo.get_project_by_id(project_id)
        if project and project.task and not tasks:
            return project.task.name
        
        return ", ".join([task.name for task in tasks]) if tasks else ""
    
    def _generate_technical_skills_data(self) -> Dict[str, List[Dict[str, Any]]]:
        """技術スキルデータを生成"""
        skills_data = {}
        
        categories = [
            ('os', 'OS'),
            ('language', '言語'),
            ('db', 'DB'),
            ('framework', 'FW/ライブラリ'),
            ('tool', 'ツール'),
            ('cloud', 'クラウド')
        ]
        
        for kind, label in categories:
            stats = self.stats_service.get_all_tech_stats(kind)
            
            skills_data[label] = [
                {
                    "name": stat["name"],
                    "experience": stat["display"],
                    "months": stat["months"]
                }
                for stat in stats
            ]
        
        return skills_data
    
    def _generate_self_pr_data(self) -> List[Dict[str, str]]:
        """自己PRデータを生成"""
        prs = self.repo.get_all_self_prs()
        
        return [
            {
                "title": pr.title,
                "content": pr.content
            }
            for pr in prs
        ]
    
    def _generate_qualification_data(self) -> List[Dict[str, str]]:
        """取得資格データを生成"""
        qualifications = self.repo.get_all_user_qualifications()
        
        return [
            {
                "name": qual.qualification.name,
                "date": qual.obtained_date.strftime("%Y年%m月") if qual.obtained_date else "",
                "note": qual.note or ""
            }
            for qual in qualifications
        ]
    
    def _generate_other_experience_data(self) -> List[Dict[str, str]]:
        """その他経歴データを生成"""
        experiences = self.repo.get_all_other_experiences()
        
        return [
            {
                "title": exp.title,
                "period": self._format_period_simple(exp.start_date, exp.end_date),
                "content": exp.content,
                "note": exp.note or ""
            }
            for exp in experiences
        ]
    
    def _format_period_simple(self, start_date, end_date) -> str:
        """シンプルな期間フォーマット"""
        if not start_date and not end_date:
            return ""
        
        try:
            start_str = ""
            end_str = ""
            
            if start_date:
                start_str = start_date.strftime("%Y年%m月")
            
            if end_date:
                end_str = end_date.strftime("%Y年%m月")
            else:
                end_str = "継続中"
            
            if start_str and end_str:
                return f"{start_str}～{end_str}"
            elif start_str:
                return f"{start_str}～"
            elif end_str:
                return f"～{end_str}"
            else:
                return ""
        except:
            return ""
    
    def _format_period(self, start_date: Optional[str], end_date: Optional[str]) -> str:
        """期間をフォーマット"""
        if not start_date:
            return ""
        
        try:
            start = datetime.strptime(start_date, "%Y-%m-%d")
            start_str = f"{start.year}年{start.month}月"
            
            if end_date:
                end = datetime.strptime(end_date, "%Y-%m-%d")
                end_str = f"{end.year}年{end.month}月"
                return f"{start_str}\n｜\n{end_str}"
            else:
                return f"{start_str}\n｜\n現在"
        except:
            return ""
    
    def _format_period_with_duration(self, start_date: Optional[str], end_date: Optional[str]) -> str:
        """期間と期間長をフォーマット"""
        period = self._format_period(start_date, end_date)
        
        if not period:
            return ""
        
        duration = self._calculate_duration(start_date, end_date)
        
        if duration:
            return f"{period}\n（{duration}）"
        
        return period
    
    def _calculate_duration(self, start_date: Optional[str], end_date: Optional[str]) -> str:
        """期間の長さを計算"""
        if not start_date:
            return ""
        
        try:
            start = datetime.strptime(start_date, "%Y-%m-%d")
            
            if end_date:
                end = datetime.strptime(end_date, "%Y-%m-%d")
            else:
                end = datetime.now()
            
            # 月数を計算
            months = (end.year - start.year) * 12 + (end.month - start.month) + 1
            
            years = months // 12
            remaining_months = months % 12
            
            if years > 0 and remaining_months > 0:
                return f"{years}年{remaining_months}ヶ月"
            elif years > 0:
                return f"{years}年"
            else:
                return f"{remaining_months}ヶ月"
        except:
            return ""
    
    def _get_career_period(self) -> str:
        """開発経歴期間を取得"""
        projects = self.repo.get_all_projects()
        if not projects:
            return "期間未設定"
        
        # 最も古い開始日と最新の終了日を取得
        start_dates = [p.project_start for p in projects if p.project_start]
        
        if not start_dates:
            return "期間未設定"
        
        earliest_start = min(start_dates)
        
        try:
            start_date = datetime.strptime(earliest_start, "%Y-%m-%d")
            return f"{start_date.year}年{start_date.month:02d}月～現在"
        except:
            return "期間未設定"
    
    def export_to_docx(self, filepath: str, name: str = "氏名"):
        """DOCXファイルとしてエクスポート
        
        Args:
            filepath: 出力ファイルパス
            name: 氏名
        """
        data = self.generate_skill_sheet_data(name)
        doc = Document()
        
        # ページ設定とマージン（狭めに調整）
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.0)    # 2.54から2.0に
            section.bottom_margin = Cm(2.0) # 2.54から2.0に
            section.left_margin = Cm(2.0)   # 2.54から2.0に
            section.right_margin = Cm(2.0)  # 2.54から2.0に
            section.page_height = Cm(29.7)  # A4
            section.page_width = Cm(21.0)
        
        # ドキュメント全体のフォント設定
        style = doc.styles['Normal']
        font = style.font
        font.name = 'MS 明朝'
        font.size = Pt(9)
        
        # ヘッダー（段落として作成、下線なし）
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(data["header"]["title"])
        title_run.font.name = 'MS 明朝'
        title_run.font.size = Pt(9)
        title_run.font.bold = True
        
        # 日付と氏名を右寄せ
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(data["header"]["date"])
        date_run.font.name = 'MS 明朝'
        date_run.font.size = Pt(9)
        
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        name_run = name_para.add_run(f"氏名：{data['header']['name']}")
        name_run.font.name = 'MS 明朝'
        name_run.font.size = Pt(9)
        
        # 開発経歴セクション
        history_para = doc.add_paragraph()
        history_run = history_para.add_run(f"開発経歴（{self._get_career_period()}）")
        history_run.font.name = 'MS 明朝'
        history_run.font.size = Pt(9)
        history_run.font.bold = True
        
        # プロジェクトテーブル
        if data["projects"]:
            table = doc.add_table(rows=1, cols=5)
            self._set_table_borders(table)
            
            # テーブル幅の設定
            table.autofit = False
            table.allow_autofit = False
            
            # ヘッダー行の設定
            header_cells = table.rows[0].cells
            self._format_header_cell(header_cells[0], "現場参画\n期間")
            self._format_header_cell(header_cells[1], "プロジェクト\n期間")
            self._format_header_cell(header_cells[2], "プロジェクト名および業務内容")
            self._format_header_cell(header_cells[3], "開発環境")
            self._format_header_cell(header_cells[4], "役割／担当／規模")
            
            # 列幅の設定（サンプルに合わせて調整）
            self._set_cell_properties(header_cells[0], width_cm=2.5)  # 現場参画期間
            self._set_cell_properties(header_cells[1], width_cm=2.5)  # プロジェクト期間
            self._set_cell_properties(header_cells[2], width_cm=8.0)  # プロジェクト名・業務内容
            self._set_cell_properties(header_cells[3], width_cm=3.5)  # 開発環境
            self._set_cell_properties(header_cells[4], width_cm=3.5)  # 役割・担当・規模
            
            # プロジェクトデータ
            for project in data["projects"]:
                row_cells = table.add_row().cells
                
                # 各セルの書式設定
                for i, cell in enumerate(row_cells):
                    widths = [2.5, 2.5, 8.0, 3.5, 3.5]
                    self._set_cell_properties(cell, width_cm=widths[i], vertical_align='top')
                
                # 期間セル（縦書き風に改行を入れる）
                row_cells[0].text = project["participation_period"]
                row_cells[1].text = project["project_period"]
                
                # プロジェクト詳細（サンプルの形式に合わせる）
                project_text = f"【プロジェクト】\n{project['project_name']}\n\n"
                if project["business_content"]:
                    project_text += f"【業務内容】\n{project['business_content']}\n\n"
                if project["project_detail"]:
                    project_text += f"【プロジェクト詳細】\n{project['project_detail']}"
                row_cells[2].text = project_text
                
                # 開発環境（改行で区切る）
                env_parts = []
                if project["environment"]["os"]:
                    env_parts.append(f"【OS】\n{chr(10).join(project['environment']['os'])}")
                if project["environment"]["language"]:
                    env_parts.append(f"【言語】\n{chr(10).join(project['environment']['language'])}")
                if project["environment"]["db"]:
                    env_parts.append(f"【DB】\n{chr(10).join(project['environment']['db'])}")
                if project["environment"]["framework"]:
                    env_parts.append(f"【FW/ライブラリ】\n{chr(10).join(project['environment']['framework'])}")
                if project["environment"]["tool"]:
                    # ツールは多いので主要なもののみ
                    main_tools = project['environment']['tool'][:8]
                    env_parts.append(f"【ツール】\n{chr(10).join(main_tools)}")
                if project["environment"]["cloud"]:
                    env_parts.append(f"【クラウド】\n{chr(10).join(project['environment']['cloud'])}")
                
                row_cells[3].text = '\n\n'.join(env_parts)
                
                # 役割/担当/規模（サンプルの形式に合わせる）
                role_parts = []
                if project["role"]:
                    role_parts.append(f"【役割】\n{project['role']}")
                if project["task"]:
                    role_parts.append(f"【担当】\n{project['task']}")
                if project["team_size"]:
                    role_parts.append(f"【プロジェクト規模】\n{project['team_size']}")
                
                row_cells[4].text = '\n\n'.join(role_parts)
        
        # テクニカルスキル
        doc.add_page_break()
        skill_heading = doc.add_paragraph()
        skill_run = skill_heading.add_run("■　テクニカルスキル")
        skill_run.font.name = 'MS 明朝'
        skill_run.font.size = Pt(9)
        skill_run.font.bold = True
        
        if data["technical_skills"]:
            skill_table = doc.add_table(rows=1, cols=4)
            self._set_table_borders(skill_table)
            
            # ヘッダー行の設定
            header_cells = skill_table.rows[0].cells
            self._format_header_cell(header_cells[0], "カテゴリ")
            self._format_header_cell(header_cells[1], "技術")
            self._format_header_cell(header_cells[2], "経験期間")
            self._format_header_cell(header_cells[3], "補足")
            
            # 列幅設定
            self._set_cell_properties(header_cells[0], width_cm=2.0)  # カテゴリ
            self._set_cell_properties(header_cells[1], width_cm=6.0)  # 技術
            self._set_cell_properties(header_cells[2], width_cm=4.0)  # 経験期間
            self._set_cell_properties(header_cells[3], width_cm=8.0)  # 補足
            
            # カテゴリ順序をサンプルに合わせる
            category_order = ["OS", "言語", "DB", "FW/ライブラリ", "ツール", "クラウド"]
            
            for category in category_order:
                if category in data["technical_skills"] and data["technical_skills"][category]:
                    skills = data["technical_skills"][category]
                    
                    # 技術名と期間を改行で区切る
                    tech_names = "\n".join([s["name"] for s in skills])
                    tech_experiences = "\n".join([s["experience"] for s in skills])
                    
                    row_cells = skill_table.add_row().cells
                    
                    # 各セルの書式設定
                    for i, cell in enumerate(row_cells):
                        widths = [2.0, 6.0, 4.0, 8.0]
                        self._set_cell_properties(cell, width_cm=widths[i], vertical_align='top')
                    
                    row_cells[0].text = category
                    row_cells[1].text = tech_names
                    row_cells[2].text = tech_experiences
                    row_cells[3].text = ""  # 補足は空欄（将来的に追加可能）
        
        # 取得資格
        if data["qualifications"]:
            qualification_heading = doc.add_paragraph()
            qualification_run = qualification_heading.add_run("■　取得資格")
            qualification_run.font.name = 'MS 明朝'
            qualification_run.font.size = Pt(9)
            qualification_run.font.bold = True
            
            qual_table = doc.add_table(rows=1, cols=3)
            self._set_table_borders(qual_table)
            
            # ヘッダー行
            qual_header_cells = qual_table.rows[0].cells
            self._format_header_cell(qual_header_cells[0], "資格名")
            self._format_header_cell(qual_header_cells[1], "取得年月")
            self._format_header_cell(qual_header_cells[2], "備考")
            
            # 列幅設定
            self._set_cell_properties(qual_header_cells[0], width_cm=8.0)
            self._set_cell_properties(qual_header_cells[1], width_cm=3.0)
            self._set_cell_properties(qual_header_cells[2], width_cm=9.0)
            
            # 資格データ
            for qual in data["qualifications"]:
                row_cells = qual_table.add_row().cells
                
                for i, cell in enumerate(row_cells):
                    widths = [8.0, 3.0, 9.0]
                    self._set_cell_properties(cell, width_cm=widths[i], vertical_align='top')
                
                row_cells[0].text = qual["name"]
                row_cells[1].text = qual["date"]
                row_cells[2].text = qual["note"]
        
        # その他経歴
        if data["other_experiences"]:
            other_heading = doc.add_paragraph()
            other_run = other_heading.add_run("■　その他経歴（学習）")
            other_run.font.name = 'MS 明朝'
            other_run.font.size = Pt(9)
            other_run.font.bold = True
            
            for exp in data["other_experiences"]:
                # タイトルと期間
                title_para = doc.add_paragraph()
                title_text = exp["title"]
                if exp["period"]:
                    title_text += f"　（{exp['period']}）"
                title_run = title_para.add_run(f"◆{title_text}")
                title_run.font.name = 'MS 明朝'
                title_run.font.size = Pt(9)
                title_run.font.bold = True
                
                # 内容
                content_para = doc.add_paragraph()
                content_run = content_para.add_run(exp["content"])
                content_run.font.name = 'MS 明朝'
                content_run.font.size = Pt(9)
                
                # 備考があれば追加
                if exp["note"]:
                    note_para = doc.add_paragraph()
                    note_run = note_para.add_run(f"※{exp['note']}")
                    note_run.font.name = 'MS 明朝'
                    note_run.font.size = Pt(9)
                    note_run.font.italic = True
                
                doc.add_paragraph()  # 項目間のスペース
        
        # 自己PR
        if data["self_pr"]:
            doc.add_page_break()
            pr_heading = doc.add_paragraph()
            pr_run = pr_heading.add_run("■　自己PR")
            pr_run.font.name = 'MS 明朝'
            pr_run.font.size = Pt(9)
            pr_run.font.bold = True
            
            for i, pr in enumerate(data["self_pr"]):
                if i > 0:
                    doc.add_paragraph()  # PR項目間にスペース
                
                # サブタイトル
                subtitle_para = doc.add_paragraph()
                subtitle_run = subtitle_para.add_run(f"◆{pr['title']}")
                subtitle_run.font.name = 'MS 明朝'
                subtitle_run.font.size = Pt(9)
                subtitle_run.font.bold = True
                
                # 内容
                content_para = doc.add_paragraph()
                content_run = content_para.add_run(pr['content'])
                content_run.font.name = 'MS 明朝'
                content_run.font.size = Pt(9)
                
                # 行間設定
                content_para.paragraph_format.line_spacing = 1.2
                content_para.paragraph_format.space_after = Pt(6)
        
        # 保存
        doc.save(filepath)
    
    def export_to_markdown(self, filepath: str, name: str = "氏名"):
        """Markdownファイルとしてエクスポート
        
        Args:
            filepath: 出力ファイルパス
            name: 氏名
        """
        data = self.generate_skill_sheet_data(name)
        
        content = []
        content.append(f"# {data['header']['title']}")
        content.append("")
        content.append(data['header']['date'])
        content.append(f"氏名：{data['header']['name']}")
        content.append("")
        
        # 開発経歴
        content.append("## 開発経歴")
        content.append("")
        
        for project in data["projects"]:
            content.append(f"### {project['project_name']}")
            content.append("")
            period = project['project_period'].replace('\n', ' ')
            content.append(f"**期間**: {period}")
            content.append(f"**契約会社**: {project['contract_company']}")
            if project['end_user']:
                content.append(f"**エンドユーザー**: {project['end_user']}")
            content.append("")
            
            if project['business_content']:
                content.append("**業務内容**")
                content.append(project['business_content'])
                content.append("")
            
            if project['project_detail']:
                content.append("**詳細**")
                content.append(project['project_detail'])
                content.append("")
            
            content.append("**開発環境**")
            for key, label in [("os", "OS"), ("language", "言語"), ("db", "DB"), 
                              ("framework", "FW"), ("tool", "ツール"), ("cloud", "クラウド")]:
                if project['environment'][key]:
                    content.append(f"- {label}: {', '.join(project['environment'][key])}")
            content.append("")
            
            if project['role'] or project['task'] or project['team_size']:
                content.append("**体制**")
                if project['role']:
                    content.append(f"- 役割: {project['role']}")
                if project['task']:
                    content.append(f"- 担当: {project['task']}")
                if project['team_size']:
                    content.append(f"- 規模: {project['team_size']}")
                content.append("")
        
        # テクニカルスキル
        content.append("## テクニカルスキル")
        content.append("")
        
        for category, skills in data["technical_skills"].items():
            if skills:
                content.append(f"### {category}")
                content.append("")
                content.append("| 技術 | 経験期間 |")
                content.append("|------|----------|")
                for skill in skills:
                    content.append(f"| {skill['name']} | {skill['experience']} |")
                content.append("")
        
        # 取得資格
        if data["qualifications"]:
            content.append("## 取得資格")
            content.append("")
            content.append("| 資格名 | 取得年月 | 備考 |")
            content.append("|--------|----------|------|")
            for qual in data["qualifications"]:
                content.append(f"| {qual['name']} | {qual['date']} | {qual['note']} |")
            content.append("")
        
        # その他経歴
        if data["other_experiences"]:
            content.append("## その他経歴（学習）")
            content.append("")
            
            for exp in data["other_experiences"]:
                title = exp['title']
                if exp['period']:
                    title += f"　（{exp['period']}）"
                content.append(f"### ◆{title}")
                content.append("")
                content.append(exp['content'])
                if exp['note']:
                    content.append("")
                    content.append(f"※{exp['note']}")
                content.append("")
        
        # 自己PR
        if data["self_pr"]:
            content.append("## 自己PR")
            content.append("")
            
            for pr in data["self_pr"]:
                content.append(f"### ◆{pr['title']}")
                content.append("")
                content.append(pr['content'])
                content.append("")
        
        # ファイルに書き込み
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("\n".join(content))